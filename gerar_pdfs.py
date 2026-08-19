#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from __future__ import annotations

import json
import shutil
import subprocess
from datetime import datetime
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent
SAIDA = ROOT / "saida"

PLACEHOLDERS = {
    "«RECLAMANTE»": "nome",
    "«CPF»": "cpf",
    "«RG»": "rg",
    "«PIS»": "pis",
    "«ENDEREÇO»": "endereco",
    "«DATA_ATUAL»": "data",
}

MODELOS = {
    "contrato": (ROOT / "modelo_contrato.docx", "contrato"),
    "procuracao": (ROOT / "modelo_procuracao.docx", "procuracao"),
    "hipossuficiencia": (ROOT / "modelo_hipossuficiencia.docx", "declaracao_hipossuficiencia"),
}


def carregar_dados(arquivo_json: str | Path = ROOT / "dados.json") -> dict:
    with open(arquivo_json, "r", encoding="utf-8") as arquivo:
        return json.load(arquivo)


def valor_dado(dados: dict, campo: str) -> str:
    valor = str(dados.get(campo, "") or "").strip()
    if valor:
        return valor
    if campo == "data":
        return datetime.now().strftime("%d/%m/%Y")
    return ""


def substituir_texto(texto: str, dados: dict) -> str:
    for placeholder, campo in PLACEHOLDERS.items():
        texto = texto.replace(placeholder, valor_dado(dados, campo))
    return texto


def substituir_em_paragrafo(paragrafo, dados: dict) -> None:
    # Primeiro tenta preservar o estilo dos runs, que é o caso normal dos modelos.
    encontrou = False
    for run in paragrafo.runs:
        novo = substituir_texto(run.text or "", dados)
        if novo != run.text:
            run.text = novo
            encontrou = True
    # Alguns editores dividem o placeholder entre vários runs; nesse caso,
    # reescrevemos apenas o parágrafo, preservando o formato do parágrafo.
    texto_atual = "".join(run.text or "" for run in paragrafo.runs)
    texto_novo = substituir_texto(texto_atual, dados)
    if texto_novo != texto_atual and not encontrou:
        paragrafo.text = texto_novo


def substituir_no_documento(documento: Document, dados: dict) -> None:
    for paragrafo in documento.paragraphs:
        substituir_em_paragrafo(paragrafo, dados)
    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    substituir_em_paragrafo(paragrafo, dados)
    for secao in documento.sections:
        for parte in (secao.header, secao.footer):
            for paragrafo in parte.paragraphs:
                substituir_em_paragrafo(paragrafo, dados)


def converter_para_pdf(arquivo_docx: Path) -> Path | None:
    arquivo_pdf = arquivo_docx.with_suffix(".pdf")
    try:
        executavel = shutil.which("libreoffice") or shutil.which("soffice")
        if executavel:
            subprocess.run(
                [executavel, "--headless", "--convert-to", "pdf", "--outdir", str(arquivo_docx.parent), str(arquivo_docx)],
                check=True,
                capture_output=True,
            )
        else:
            from docx2pdf import convert
            convert(str(arquivo_docx), str(arquivo_pdf))
        return arquivo_pdf if arquivo_pdf.exists() else None
    except Exception as erro:
        print(f"Aviso: não foi possível criar {arquivo_pdf.name}: {erro}")
        return None


def gerar_documentos_a_partir_dados(dados_cliente: dict) -> list[str]:
    SAIDA.mkdir(exist_ok=True)
    nome = valor_dado(dados_cliente, "nome") or "cliente"
    nome_seguro = "_".join(parte for parte in nome.lower().split() if parte.isalnum())
    arquivos_gerados: list[str] = []
    for tipo, (modelo, prefixo) in MODELOS.items():
        if not modelo.exists():
            raise FileNotFoundError(f"Modelo ausente: {modelo.name}")
        destino_docx = SAIDA / f"{prefixo}_{nome_seguro}.docx"
        documento = Document(str(modelo))
        substituir_no_documento(documento, dados_cliente)
        documento.save(str(destino_docx))
        arquivos_gerados.append(str(destino_docx))
        pdf = converter_para_pdf(destino_docx)
        if pdf:
            arquivos_gerados.append(str(pdf))
    return arquivos_gerados


def gerar_documentos() -> list[str]:
    dados = carregar_dados().get("cliente", {})
    arquivos = gerar_documentos_a_partir_dados(dados)
    print(f"Processo concluído: {len(arquivos)} arquivo(s) em {SAIDA}")
    return arquivos


if __name__ == "__main__":
    gerar_documentos()
